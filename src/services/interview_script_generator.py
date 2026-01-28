"""
Interview Script Document Generator Service

Generates interview scripts in multiple formats (PDF, DOCX, Markdown)
and stores them in project-specific directories.
"""

import os
import logging
from datetime import datetime
from typing import Dict, Any, Optional, List
from pathlib import Path

from config.settings import settings

logger = logging.getLogger(__name__)


class InterviewScriptGenerator:
    """
    Service for generating interview script documents in multiple formats.

    Supports:
    - PDF: Professional formatted document using ReportLab
    - DOCX: Editable Word document using python-docx
    - Markdown: Plain text format for easy viewing/editing
    """

    def __init__(self):
        """Initialize the generator with configured scripts directory."""
        self.scripts_dir = Path(settings.SCRIPTS_DIR)
        self.scripts_dir.mkdir(parents=True, exist_ok=True)

    def get_project_scripts_dir(self, project_id: str) -> Path:
        """
        Get or create the scripts directory for a specific project.

        Args:
            project_id: The project ID

        Returns:
            Path to the project's scripts directory
        """
        project_dir = self.scripts_dir / project_id
        project_dir.mkdir(parents=True, exist_ok=True)
        return project_dir

    def generate_all_formats(
        self,
        interview_script: Dict[str, Any],
        project: Dict[str, Any],
    ) -> Dict[str, str]:
        """
        Generate interview script in all supported formats.

        Args:
            interview_script: The interview script data
            project: Project information

        Returns:
            Dictionary mapping format names to file paths
        """
        project_id = interview_script.get("project_id", project.get("id", "unknown"))

        paths = {}

        try:
            paths["pdf"] = self.generate_pdf(interview_script, project)
            logger.info(f"Generated PDF interview script: {paths['pdf']}")
        except Exception as e:
            logger.error(f"Failed to generate PDF: {e}")

        try:
            paths["docx"] = self.generate_docx(interview_script, project)
            logger.info(f"Generated DOCX interview script: {paths['docx']}")
        except Exception as e:
            logger.error(f"Failed to generate DOCX: {e}")

        try:
            paths["markdown"] = self.generate_markdown(interview_script, project)
            logger.info(f"Generated Markdown interview script: {paths['markdown']}")
        except Exception as e:
            logger.error(f"Failed to generate Markdown: {e}")

        return paths

    def generate_pdf(
        self,
        interview_script: Dict[str, Any],
        project: Dict[str, Any],
    ) -> str:
        """
        Generate interview script as a PDF document.

        Args:
            interview_script: The interview script data
            project: Project information

        Returns:
            Path to the generated PDF file
        """
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            PageBreak, ListFlowable, ListItem
        )
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT

        project_id = interview_script.get("project_id", project.get("id", "unknown"))
        project_dir = self.get_project_scripts_dir(project_id)

        # Generate filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        client_name = project.get("client_name", "client").replace(" ", "_")
        filename = f"interview_script_{client_name}_{timestamp}.pdf"
        filepath = project_dir / filename

        # Create document
        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72,
        )

        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=TA_CENTER,
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            spaceAfter=12,
            spaceBefore=20,
        )
        subheading_style = ParagraphStyle(
            'CustomSubHeading',
            parent=styles['Heading3'],
            fontSize=12,
            spaceAfter=8,
            spaceBefore=12,
        )
        question_style = ParagraphStyle(
            'QuestionStyle',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            leftIndent=20,
        )
        intent_style = ParagraphStyle(
            'IntentStyle',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.grey,
            leftIndent=20,
            spaceAfter=4,
        )

        # Build content
        content = []

        # Title
        content.append(Paragraph("Interview Script", title_style))
        content.append(Paragraph(
            f"Project: {project.get('project_name', 'N/A')}",
            ParagraphStyle('SubTitle', parent=styles['Normal'], alignment=TA_CENTER)
        ))
        content.append(Paragraph(
            f"Client: {project.get('client_name', 'N/A')}",
            ParagraphStyle('SubTitle', parent=styles['Normal'], alignment=TA_CENTER)
        ))
        content.append(Spacer(1, 12))

        # Metadata table
        generated_at = interview_script.get("generated_at", datetime.now().isoformat())
        if isinstance(generated_at, str):
            try:
                generated_at = datetime.fromisoformat(generated_at.replace("Z", "+00:00"))
                generated_at = generated_at.strftime("%B %d, %Y at %I:%M %p")
            except Exception:
                pass

        meta_data = [
            ["Target Departments", ", ".join(interview_script.get("target_departments", []))],
            ["Target Roles", ", ".join(interview_script.get("target_roles", []))],
            ["Estimated Duration", f"{interview_script.get('estimated_duration_minutes', 60)} minutes"],
            ["Generated", str(generated_at)],
        ]
        meta_table = Table(meta_data, colWidths=[2*inch, 4*inch])
        meta_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('PADDING', (0, 0), (-1, -1), 8),
        ]))
        content.append(meta_table)
        content.append(Spacer(1, 24))

        # Purpose section
        content.append(Paragraph("Purpose", heading_style))
        content.append(Paragraph(
            "This interview script is designed to help identify improvement opportunities "
            "and potential cost savings in the organization's processes. The questions focus on "
            "uncovering inefficiencies, workarounds, and areas where automation could provide value.",
            styles['Normal']
        ))
        content.append(Spacer(1, 12))

        # Customer Context Section
        customer_context = interview_script.get("customer_context", {})
        if customer_context:
            content.append(Paragraph("Customer Context", heading_style))

            if customer_context.get("business_overview"):
                content.append(Paragraph("<b>Business Overview:</b>", styles['Normal']))
                content.append(Paragraph(customer_context.get("business_overview", ""), styles['Normal']))
                content.append(Spacer(1, 6))

            if customer_context.get("organization_structure"):
                content.append(Paragraph("<b>Organization Structure:</b>", styles['Normal']))
                content.append(Paragraph(customer_context.get("organization_structure", ""), styles['Normal']))
                content.append(Spacer(1, 6))

            if customer_context.get("industry_context"):
                content.append(Paragraph("<b>Industry Context:</b>", styles['Normal']))
                content.append(Paragraph(customer_context.get("industry_context", ""), styles['Normal']))
                content.append(Spacer(1, 6))

            current_challenges = customer_context.get("current_challenges", [])
            if current_challenges:
                content.append(Paragraph("<b>Current Challenges Identified:</b>", styles['Normal']))
                for challenge in current_challenges:
                    content.append(Paragraph(f"  • {challenge}", styles['Normal']))
                content.append(Spacer(1, 6))

            key_processes = customer_context.get("key_processes", [])
            if key_processes:
                content.append(Paragraph("<b>Key Processes Under Review:</b>", styles['Normal']))
                for process in key_processes:
                    content.append(Paragraph(f"  • {process}", styles['Normal']))
                content.append(Spacer(1, 6))

            stakeholders = customer_context.get("stakeholders", [])
            if stakeholders:
                content.append(Paragraph("<b>Key Stakeholders:</b>", styles['Normal']))
                content.append(Paragraph(", ".join(stakeholders), styles['Normal']))
                content.append(Spacer(1, 6))

            if customer_context.get("data_sources_summary"):
                content.append(Paragraph("<b>Data Sources Analyzed:</b>", styles['Normal']))
                content.append(Paragraph(customer_context.get("data_sources_summary", ""), styles['Normal']))

            content.append(Spacer(1, 12))

        # Introduction
        content.append(Paragraph("Introduction", heading_style))
        introduction = interview_script.get("introduction", "")
        for para in introduction.split("\n\n"):
            if para.strip():
                content.append(Paragraph(para.strip(), styles['Normal']))
                content.append(Spacer(1, 6))

        content.append(PageBreak())

        # Diagnostic Form - Leads to Validate
        diagnostic_leads = interview_script.get("diagnostic_leads", [])
        if diagnostic_leads:
            content.append(Paragraph("Diagnostic Form: AI-Generated Leads to Validate", heading_style))
            content.append(Paragraph(
                "The following leads were identified through AI analysis of the customer's data. "
                "Use this form to validate or invalidate each lead during interviews.",
                styles['Normal']
            ))
            content.append(Spacer(1, 12))

            lead_num = 1
            for lead in diagnostic_leads:
                # Lead header with confidence
                confidence = lead.get("confidence", 0.5)
                confidence_pct = int(confidence * 100)
                content.append(Paragraph(
                    f"<b>Lead #{lead_num}:</b> {lead.get('lead_summary', '')} "
                    f"<i>(Confidence: {confidence_pct}%)</i>",
                    subheading_style
                ))

                # Category
                content.append(Paragraph(
                    f"Category: {lead.get('category', 'General')}",
                    intent_style
                ))

                # Validation questions
                validation_questions = lead.get("validation_questions", [])
                if validation_questions:
                    content.append(Paragraph("Validation Questions:", styles['Normal']))
                    for vq in validation_questions:
                        content.append(Paragraph(f"    □ {vq}", question_style))

                # Expected evidence
                expected_evidence = lead.get("expected_evidence", [])
                if expected_evidence:
                    content.append(Spacer(1, 6))
                    content.append(Paragraph("Expected Evidence (if lead is valid):", intent_style))
                    for ev in expected_evidence:
                        content.append(Paragraph(f"    • {ev}", intent_style))

                # Validation checkbox
                content.append(Spacer(1, 6))
                content.append(Paragraph(
                    "Validation Result:  □ Confirmed  □ Partially Confirmed  □ Not Confirmed  □ Needs More Info",
                    styles['Normal']
                ))
                content.append(Paragraph("Notes: _" * 40, styles['Normal']))
                content.append(Spacer(1, 12))
                lead_num += 1

            content.append(PageBreak())

        # Questions by Role - Validation Questions
        content.append(Paragraph("Part A: Validation Questions", heading_style))
        content.append(Paragraph(
            "These questions are designed to validate or invalidate the AI-generated leads above.",
            styles['Normal']
        ))
        content.append(Spacer(1, 12))

        questions = interview_script.get("questions", [])

        # Group questions by role
        questions_by_role: Dict[str, List] = {}
        for q in questions:
            role = q.get("role", "General")
            if role not in questions_by_role:
                questions_by_role[role] = []
            questions_by_role[role].append(q)

        question_num = 1
        for role, role_questions in questions_by_role.items():
            content.append(Paragraph(f"Questions for: {role}", subheading_style))

            for q in role_questions:
                # Main question
                content.append(Paragraph(
                    f"<b>Q{question_num}:</b> {q.get('question', '')}",
                    question_style
                ))

                # Intent
                intent = q.get("intent", "")
                if intent:
                    content.append(Paragraph(
                        f"<i>Intent: {intent}</i>",
                        intent_style
                    ))

                # Follow-ups
                follow_ups = q.get("follow_ups", [])
                if follow_ups:
                    content.append(Paragraph("Follow-up questions:", intent_style))
                    for fu in follow_ups:
                        content.append(Paragraph(f"    - {fu}", intent_style))

                content.append(Spacer(1, 12))
                question_num += 1

            content.append(Spacer(1, 12))

        # Part B: Discovery Questions
        discovery_questions = interview_script.get("discovery_questions", [])
        if discovery_questions:
            content.append(PageBreak())
            content.append(Paragraph("Part B: Discovery Questions", heading_style))
            content.append(Paragraph(
                "These open-ended questions are designed to identify NEW improvement opportunities "
                "specific to this customer that may not have been captured in the data analysis.",
                styles['Normal']
            ))
            content.append(Spacer(1, 12))

            # Group discovery questions by role
            disc_by_role: Dict[str, List] = {}
            for q in discovery_questions:
                role = q.get("role", "General")
                if role not in disc_by_role:
                    disc_by_role[role] = []
                disc_by_role[role].append(q)

            disc_num = 1
            for role, role_questions in disc_by_role.items():
                content.append(Paragraph(f"Discovery Questions for: {role}", subheading_style))

                for q in role_questions:
                    content.append(Paragraph(
                        f"<b>D{disc_num}:</b> {q.get('question', '')}",
                        question_style
                    ))

                    intent = q.get("intent", "")
                    if intent:
                        content.append(Paragraph(
                            f"<i>Intent: {intent}</i>",
                            intent_style
                        ))

                    follow_ups = q.get("follow_ups", [])
                    if follow_ups:
                        content.append(Paragraph("Follow-up questions:", intent_style))
                        for fu in follow_ups:
                            content.append(Paragraph(f"    - {fu}", intent_style))

                    content.append(Spacer(1, 12))
                    disc_num += 1

        # Closing Notes
        content.append(PageBreak())
        content.append(Paragraph("Closing Notes", heading_style))
        closing_notes = interview_script.get("closing_notes", "")
        for para in closing_notes.split("\n\n"):
            if para.strip():
                content.append(Paragraph(para.strip(), styles['Normal']))
                content.append(Spacer(1, 6))

        # Tips for Interviewers
        content.append(Spacer(1, 24))
        content.append(Paragraph("Tips for Interviewers", heading_style))
        tips = [
            "Listen actively and take detailed notes on specific examples provided.",
            "Look for quantifiable impacts - hours spent, frequency of issues, number of people affected.",
            "Ask for concrete examples when interviewees describe problems.",
            "Note any workarounds or unofficial processes mentioned.",
            "Pay attention to emotional responses - frustration often indicates significant pain points.",
            "Ask 'why' multiple times to get to root causes.",
        ]
        for tip in tips:
            content.append(Paragraph(f"- {tip}", styles['Normal']))

        # Build PDF
        doc.build(content)

        logger.info(f"PDF interview script generated: {filepath}")
        return str(filepath)

    def generate_docx(
        self,
        interview_script: Dict[str, Any],
        project: Dict[str, Any],
    ) -> str:
        """
        Generate interview script as a Word document.

        Args:
            interview_script: The interview script data
            project: Project information

        Returns:
            Path to the generated DOCX file
        """
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE

        project_id = interview_script.get("project_id", project.get("id", "unknown"))
        project_dir = self.get_project_scripts_dir(project_id)

        # Generate filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        client_name = project.get("client_name", "client").replace(" ", "_")
        filename = f"interview_script_{client_name}_{timestamp}.docx"
        filepath = project_dir / filename

        # Create document
        doc = Document()

        # Title
        title = doc.add_heading("Interview Script", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"Project: {project.get('project_name', 'N/A')}")
        run.bold = True
        subtitle.add_run(f"\nClient: {project.get('client_name', 'N/A')}")

        # Metadata
        doc.add_paragraph()
        meta_table = doc.add_table(rows=4, cols=2)
        meta_table.style = 'Table Grid'

        meta_data = [
            ("Target Departments", ", ".join(interview_script.get("target_departments", []))),
            ("Target Roles", ", ".join(interview_script.get("target_roles", []))),
            ("Estimated Duration", f"{interview_script.get('estimated_duration_minutes', 60)} minutes"),
            ("Generated", datetime.now().strftime("%B %d, %Y at %I:%M %p")),
        ]

        for i, (label, value) in enumerate(meta_data):
            row = meta_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            row.cells[0].paragraphs[0].runs[0].bold = True

        # Purpose
        doc.add_heading("Purpose", level=1)
        doc.add_paragraph(
            "This interview script is designed to help identify improvement opportunities "
            "and potential cost savings in the organization's processes. The questions focus on "
            "uncovering inefficiencies, workarounds, and areas where automation could provide value."
        )

        # Customer Context Section
        customer_context = interview_script.get("customer_context", {})
        if customer_context:
            doc.add_heading("Customer Context", level=1)

            if customer_context.get("business_overview"):
                p = doc.add_paragraph()
                p.add_run("Business Overview: ").bold = True
                p.add_run(customer_context.get("business_overview", ""))

            if customer_context.get("organization_structure"):
                p = doc.add_paragraph()
                p.add_run("Organization Structure: ").bold = True
                p.add_run(customer_context.get("organization_structure", ""))

            if customer_context.get("industry_context"):
                p = doc.add_paragraph()
                p.add_run("Industry Context: ").bold = True
                p.add_run(customer_context.get("industry_context", ""))

            current_challenges = customer_context.get("current_challenges", [])
            if current_challenges:
                p = doc.add_paragraph()
                p.add_run("Current Challenges Identified:").bold = True
                for challenge in current_challenges:
                    doc.add_paragraph(challenge, style='List Bullet')

            key_processes = customer_context.get("key_processes", [])
            if key_processes:
                p = doc.add_paragraph()
                p.add_run("Key Processes Under Review:").bold = True
                for process in key_processes:
                    doc.add_paragraph(process, style='List Bullet')

            stakeholders = customer_context.get("stakeholders", [])
            if stakeholders:
                p = doc.add_paragraph()
                p.add_run("Key Stakeholders: ").bold = True
                p.add_run(", ".join(stakeholders))

            if customer_context.get("data_sources_summary"):
                p = doc.add_paragraph()
                p.add_run("Data Sources Analyzed: ").bold = True
                p.add_run(customer_context.get("data_sources_summary", ""))

        # Introduction
        doc.add_heading("Introduction", level=1)
        introduction = interview_script.get("introduction", "")
        for para in introduction.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

        # Diagnostic Form - Leads to Validate
        diagnostic_leads = interview_script.get("diagnostic_leads", [])
        if diagnostic_leads:
            doc.add_page_break()
            doc.add_heading("Diagnostic Form: AI-Generated Leads to Validate", level=1)
            doc.add_paragraph(
                "The following leads were identified through AI analysis of the customer's data. "
                "Use this form to validate or invalidate each lead during interviews."
            )

            lead_num = 1
            for lead in diagnostic_leads:
                confidence = lead.get("confidence", 0.5)
                confidence_pct = int(confidence * 100)

                doc.add_heading(f"Lead #{lead_num}: {lead.get('lead_summary', '')[:80]}...", level=2)

                p = doc.add_paragraph()
                p.add_run(f"Confidence: {confidence_pct}% | Category: {lead.get('category', 'General')}")
                p.runs[0].italic = True

                validation_questions = lead.get("validation_questions", [])
                if validation_questions:
                    p = doc.add_paragraph()
                    p.add_run("Validation Questions:").bold = True
                    for vq in validation_questions:
                        doc.add_paragraph(f"□ {vq}", style='List Bullet')

                expected_evidence = lead.get("expected_evidence", [])
                if expected_evidence:
                    p = doc.add_paragraph()
                    p.add_run("Expected Evidence (if lead is valid):").bold = True
                    for ev in expected_evidence:
                        doc.add_paragraph(ev, style='List Bullet')

                p = doc.add_paragraph()
                p.add_run("Validation Result: ").bold = True
                p.add_run("□ Confirmed  □ Partially Confirmed  □ Not Confirmed  □ Needs More Info")

                doc.add_paragraph("Notes: " + "_" * 60)
                doc.add_paragraph()
                lead_num += 1

        # Questions - Part A: Validation Questions
        doc.add_page_break()
        doc.add_heading("Part A: Validation Questions", level=1)
        doc.add_paragraph(
            "These questions are designed to validate or invalidate the AI-generated leads above."
        )

        questions = interview_script.get("questions", [])

        # Group questions by role
        questions_by_role: Dict[str, List] = {}
        for q in questions:
            role = q.get("role", "General")
            if role not in questions_by_role:
                questions_by_role[role] = []
            questions_by_role[role].append(q)

        question_num = 1
        for role, role_questions in questions_by_role.items():
            doc.add_heading(f"Questions for: {role}", level=2)

            for q in role_questions:
                # Main question
                p = doc.add_paragraph()
                p.add_run(f"Q{question_num}: ").bold = True
                p.add_run(q.get("question", ""))

                # Intent
                intent = q.get("intent", "")
                if intent:
                    intent_p = doc.add_paragraph()
                    intent_run = intent_p.add_run(f"Intent: {intent}")
                    intent_run.italic = True
                    intent_run.font.size = Pt(10)

                # Follow-ups
                follow_ups = q.get("follow_ups", [])
                if follow_ups:
                    fu_p = doc.add_paragraph("Follow-up questions:")
                    fu_p.runs[0].font.size = Pt(10)
                    for fu in follow_ups:
                        fu_item = doc.add_paragraph(fu, style='List Bullet')
                        fu_item.paragraph_format.left_indent = Inches(0.5)

                doc.add_paragraph()
                question_num += 1

        # Part B: Discovery Questions
        discovery_questions = interview_script.get("discovery_questions", [])
        if discovery_questions:
            doc.add_page_break()
            doc.add_heading("Part B: Discovery Questions", level=1)
            doc.add_paragraph(
                "These open-ended questions are designed to identify NEW improvement opportunities "
                "specific to this customer that may not have been captured in the data analysis."
            )

            # Group discovery questions by role
            disc_by_role: Dict[str, List] = {}
            for q in discovery_questions:
                role = q.get("role", "General")
                if role not in disc_by_role:
                    disc_by_role[role] = []
                disc_by_role[role].append(q)

            disc_num = 1
            for role, role_questions in disc_by_role.items():
                doc.add_heading(f"Discovery Questions for: {role}", level=2)

                for q in role_questions:
                    p = doc.add_paragraph()
                    p.add_run(f"D{disc_num}: ").bold = True
                    p.add_run(q.get("question", ""))

                    intent = q.get("intent", "")
                    if intent:
                        intent_p = doc.add_paragraph()
                        intent_run = intent_p.add_run(f"Intent: {intent}")
                        intent_run.italic = True
                        intent_run.font.size = Pt(10)

                    follow_ups = q.get("follow_ups", [])
                    if follow_ups:
                        fu_p = doc.add_paragraph("Follow-up questions:")
                        fu_p.runs[0].font.size = Pt(10)
                        for fu in follow_ups:
                            fu_item = doc.add_paragraph(fu, style='List Bullet')
                            fu_item.paragraph_format.left_indent = Inches(0.5)

                    doc.add_paragraph()
                    disc_num += 1

        # Closing Notes
        doc.add_page_break()
        doc.add_heading("Closing Notes", level=1)
        closing_notes = interview_script.get("closing_notes", "")
        for para in closing_notes.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

        # Tips
        doc.add_heading("Tips for Interviewers", level=1)
        tips = [
            "Listen actively and take detailed notes on specific examples provided.",
            "Look for quantifiable impacts - hours spent, frequency of issues, number of people affected.",
            "Ask for concrete examples when interviewees describe problems.",
            "Note any workarounds or unofficial processes mentioned.",
            "Pay attention to emotional responses - frustration often indicates significant pain points.",
            "Ask 'why' multiple times to get to root causes.",
        ]
        for tip in tips:
            doc.add_paragraph(tip, style='List Bullet')

        # Notes section
        doc.add_page_break()
        doc.add_heading("Interview Notes", level=1)
        doc.add_paragraph("Use this section to capture notes during the interview:")
        doc.add_paragraph()

        # Add some blank lines for notes
        for _ in range(20):
            p = doc.add_paragraph("_" * 80)
            p.paragraph_format.space_after = Pt(12)

        # Save document
        doc.save(str(filepath))

        logger.info(f"DOCX interview script generated: {filepath}")
        return str(filepath)

    def generate_markdown(
        self,
        interview_script: Dict[str, Any],
        project: Dict[str, Any],
    ) -> str:
        """
        Generate interview script as a Markdown document.

        Args:
            interview_script: The interview script data
            project: Project information

        Returns:
            Path to the generated Markdown file
        """
        project_id = interview_script.get("project_id", project.get("id", "unknown"))
        project_dir = self.get_project_scripts_dir(project_id)

        # Generate filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        client_name = project.get("client_name", "client").replace(" ", "_")
        filename = f"interview_script_{client_name}_{timestamp}.md"
        filepath = project_dir / filename

        # Build markdown content
        lines = []

        # Header
        lines.append("# Interview Script")
        lines.append("")
        lines.append(f"**Project:** {project.get('project_name', 'N/A')}")
        lines.append(f"**Client:** {project.get('client_name', 'N/A')}")
        lines.append("")

        # Metadata
        lines.append("## Overview")
        lines.append("")
        lines.append(f"| Field | Value |")
        lines.append(f"|-------|-------|")
        lines.append(f"| Target Departments | {', '.join(interview_script.get('target_departments', []))} |")
        lines.append(f"| Target Roles | {', '.join(interview_script.get('target_roles', []))} |")
        lines.append(f"| Estimated Duration | {interview_script.get('estimated_duration_minutes', 60)} minutes |")
        lines.append(f"| Generated | {datetime.now().strftime('%B %d, %Y at %I:%M %p')} |")
        lines.append("")

        # Purpose
        lines.append("## Purpose")
        lines.append("")
        lines.append("This interview script is designed to help identify improvement opportunities "
                    "and potential cost savings in the organization's processes. The questions focus on "
                    "uncovering inefficiencies, workarounds, and areas where automation could provide value.")
        lines.append("")

        # Customer Context Section
        customer_context = interview_script.get("customer_context", {})
        if customer_context:
            lines.append("## Customer Context")
            lines.append("")

            if customer_context.get("business_overview"):
                lines.append("**Business Overview:**")
                lines.append(customer_context.get("business_overview", ""))
                lines.append("")

            if customer_context.get("organization_structure"):
                lines.append("**Organization Structure:**")
                lines.append(customer_context.get("organization_structure", ""))
                lines.append("")

            if customer_context.get("industry_context"):
                lines.append("**Industry Context:**")
                lines.append(customer_context.get("industry_context", ""))
                lines.append("")

            current_challenges = customer_context.get("current_challenges", [])
            if current_challenges:
                lines.append("**Current Challenges Identified:**")
                for challenge in current_challenges:
                    lines.append(f"- {challenge}")
                lines.append("")

            key_processes = customer_context.get("key_processes", [])
            if key_processes:
                lines.append("**Key Processes Under Review:**")
                for process in key_processes:
                    lines.append(f"- {process}")
                lines.append("")

            stakeholders = customer_context.get("stakeholders", [])
            if stakeholders:
                lines.append(f"**Key Stakeholders:** {', '.join(stakeholders)}")
                lines.append("")

            if customer_context.get("data_sources_summary"):
                lines.append("**Data Sources Analyzed:**")
                lines.append(customer_context.get("data_sources_summary", ""))
                lines.append("")

        # Introduction
        lines.append("## Introduction")
        lines.append("")
        introduction = interview_script.get("introduction", "")
        lines.append(introduction)
        lines.append("")

        # Diagnostic Form - Leads to Validate
        diagnostic_leads = interview_script.get("diagnostic_leads", [])
        if diagnostic_leads:
            lines.append("---")
            lines.append("")
            lines.append("## Diagnostic Form: AI-Generated Leads to Validate")
            lines.append("")
            lines.append("The following leads were identified through AI analysis of the customer's data. "
                        "Use this form to validate or invalidate each lead during interviews.")
            lines.append("")

            lead_num = 1
            for lead in diagnostic_leads:
                confidence = lead.get("confidence", 0.5)
                confidence_pct = int(confidence * 100)

                lines.append(f"### Lead #{lead_num}: {lead.get('lead_summary', '')}")
                lines.append("")
                lines.append(f"*Confidence: {confidence_pct}% | Category: {lead.get('category', 'General')}*")
                lines.append("")

                validation_questions = lead.get("validation_questions", [])
                if validation_questions:
                    lines.append("**Validation Questions:**")
                    for vq in validation_questions:
                        lines.append(f"- [ ] {vq}")
                    lines.append("")

                expected_evidence = lead.get("expected_evidence", [])
                if expected_evidence:
                    lines.append("**Expected Evidence (if lead is valid):**")
                    for ev in expected_evidence:
                        lines.append(f"- {ev}")
                    lines.append("")

                lines.append("**Validation Result:**")
                lines.append("- [ ] Confirmed")
                lines.append("- [ ] Partially Confirmed")
                lines.append("- [ ] Not Confirmed")
                lines.append("- [ ] Needs More Info")
                lines.append("")
                lines.append("**Notes:** _______________________________________________")
                lines.append("")
                lines.append("---")
                lines.append("")
                lead_num += 1

        # Part A: Validation Questions
        lines.append("---")
        lines.append("")
        lines.append("## Part A: Validation Questions")
        lines.append("")
        lines.append("These questions are designed to validate or invalidate the AI-generated leads above.")
        lines.append("")

        questions = interview_script.get("questions", [])

        # Group questions by role
        questions_by_role: Dict[str, List] = {}
        for q in questions:
            role = q.get("role", "General")
            if role not in questions_by_role:
                questions_by_role[role] = []
            questions_by_role[role].append(q)

        question_num = 1
        for role, role_questions in questions_by_role.items():
            lines.append(f"### Questions for: {role}")
            lines.append("")

            for q in role_questions:
                # Main question
                lines.append(f"**Q{question_num}: {q.get('question', '')}**")
                lines.append("")

                # Intent
                intent = q.get("intent", "")
                if intent:
                    lines.append(f"*Intent: {intent}*")
                    lines.append("")

                # Follow-ups
                follow_ups = q.get("follow_ups", [])
                if follow_ups:
                    lines.append("Follow-up questions:")
                    for fu in follow_ups:
                        lines.append(f"- {fu}")
                    lines.append("")

                lines.append("---")
                lines.append("")
                question_num += 1

        # Part B: Discovery Questions
        discovery_questions = interview_script.get("discovery_questions", [])
        if discovery_questions:
            lines.append("## Part B: Discovery Questions")
            lines.append("")
            lines.append("These open-ended questions are designed to identify NEW improvement opportunities "
                        "specific to this customer that may not have been captured in the data analysis.")
            lines.append("")

            # Group discovery questions by role
            disc_by_role: Dict[str, List] = {}
            for q in discovery_questions:
                role = q.get("role", "General")
                if role not in disc_by_role:
                    disc_by_role[role] = []
                disc_by_role[role].append(q)

            disc_num = 1
            for role, role_questions in disc_by_role.items():
                lines.append(f"### Discovery Questions for: {role}")
                lines.append("")

                for q in role_questions:
                    lines.append(f"**D{disc_num}: {q.get('question', '')}**")
                    lines.append("")

                    intent = q.get("intent", "")
                    if intent:
                        lines.append(f"*Intent: {intent}*")
                        lines.append("")

                    follow_ups = q.get("follow_ups", [])
                    if follow_ups:
                        lines.append("Follow-up questions:")
                        for fu in follow_ups:
                            lines.append(f"- {fu}")
                        lines.append("")

                    lines.append("---")
                    lines.append("")
                    disc_num += 1

        # Closing Notes
        lines.append("## Closing Notes")
        lines.append("")
        closing_notes = interview_script.get("closing_notes", "")
        lines.append(closing_notes)
        lines.append("")

        # Tips
        lines.append("## Tips for Interviewers")
        lines.append("")
        tips = [
            "Listen actively and take detailed notes on specific examples provided.",
            "Look for quantifiable impacts - hours spent, frequency of issues, number of people affected.",
            "Ask for concrete examples when interviewees describe problems.",
            "Note any workarounds or unofficial processes mentioned.",
            "Pay attention to emotional responses - frustration often indicates significant pain points.",
            "Ask 'why' multiple times to get to root causes.",
        ]
        for tip in tips:
            lines.append(f"- {tip}")
        lines.append("")

        # Notes section
        lines.append("---")
        lines.append("")
        lines.append("## Interview Notes")
        lines.append("")
        lines.append("*Use this section to capture notes during the interview:*")
        lines.append("")
        lines.append("### Key Pain Points Identified")
        lines.append("")
        lines.append("1. ")
        lines.append("2. ")
        lines.append("3. ")
        lines.append("")
        lines.append("### Potential Improvements")
        lines.append("")
        lines.append("1. ")
        lines.append("2. ")
        lines.append("3. ")
        lines.append("")
        lines.append("### Cost/Time Savings Opportunities")
        lines.append("")
        lines.append("| Area | Current Time/Cost | Potential Savings |")
        lines.append("|------|-------------------|-------------------|")
        lines.append("|      |                   |                   |")
        lines.append("|      |                   |                   |")
        lines.append("|      |                   |                   |")
        lines.append("")

        # Write to file
        content = "\n".join(lines)
        filepath.write_text(content, encoding="utf-8")

        logger.info(f"Markdown interview script generated: {filepath}")
        return str(filepath)

    def get_script_files(self, project_id: str) -> Dict[str, Optional[str]]:
        """
        Get paths to existing interview script files for a project.

        Args:
            project_id: The project ID

        Returns:
            Dictionary mapping format names to file paths (or None if not found)
        """
        project_dir = self.get_project_scripts_dir(project_id)

        result = {
            "pdf": None,
            "docx": None,
            "markdown": None,
        }

        # Find the most recent file of each type
        for ext, key in [(".pdf", "pdf"), (".docx", "docx"), (".md", "markdown")]:
            files = list(project_dir.glob(f"*{ext}"))
            if files:
                # Get most recent file
                most_recent = max(files, key=lambda f: f.stat().st_mtime)
                result[key] = str(most_recent)

        return result

    def list_all_scripts(self, project_id: str) -> List[Dict[str, Any]]:
        """
        List all interview script files for a project.

        Args:
            project_id: The project ID

        Returns:
            List of script file metadata
        """
        project_dir = self.get_project_scripts_dir(project_id)

        scripts = []

        for file_path in project_dir.iterdir():
            if file_path.is_file() and file_path.suffix in [".pdf", ".docx", ".md"]:
                stat = file_path.stat()
                scripts.append({
                    "filename": file_path.name,
                    "path": str(file_path),
                    "format": file_path.suffix[1:],  # Remove the dot
                    "size_bytes": stat.st_size,
                    "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                    "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                })

        # Sort by modification time, most recent first
        scripts.sort(key=lambda x: x["modified_at"], reverse=True)

        return scripts


# Global instance
interview_script_generator = InterviewScriptGenerator()


def get_interview_script_generator() -> InterviewScriptGenerator:
    """Get the global interview script generator instance."""
    return interview_script_generator
